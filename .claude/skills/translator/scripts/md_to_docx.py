# -*- coding: utf-8 -*-
"""
md_to_docx.py - 通用 Markdown 转 Word 文档工具
用法: python md_to_docx.py <input.md> [output.docx]

排版规格：
- A4纸张，标准页边距
- 正文：宋体12pt + Times New Roman，1.5倍行距，首行缩进
- H1：黑体22pt加粗居中（书名/文档标题）
- H2：黑体16pt加粗（部/大节）
- H3：黑体14pt加粗（章/小节）
- 引用块：楷体10.5pt灰色（译者说明等）
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx: pip install python-docx")
    sys.exit(1)


def set_run_font(run, name_cn='宋体', name_en='Times New Roman', size=12, bold=False, color=None):
    """设置 run 的字体"""
    run.font.size = Pt(size)
    run.font.name = name_en
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def create_word_doc(md_path, docx_path):
    """将 Markdown 文件转换为 Word 文档"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 页面设置 - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 默认段落样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 跳过水平线
        if stripped == '---':
            i += 1
            continue

        # 引用块（译者说明等）
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_run_font(run, '楷体', 'Times New Roman', 10.5, color=(100, 100, 100))
            i += 1
            continue

        # H1 - 书名/文档标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_run_font(run, '黑体', 'Times New Roman', 22, bold=True)
            i += 1
            continue

        # H2 - 部/大节
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_run_font(run, '黑体', 'Times New Roman', 16, bold=True)
            i += 1
            continue

        # H3 - 章/小节
        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_run_font(run, '黑体', 'Times New Roman', 14, bold=True)
            i += 1
            continue

        # 加粗行（如副标题）
        if stripped.startswith('**') and stripped.endswith('**'):
            text = stripped[2:-2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_run_font(run, '宋体', 'Times New Roman', 12, bold=True)
            i += 1
            continue

        # 斜体行（如作者信息）
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            text = stripped[1:-1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            set_run_font(run, '楷体', 'Times New Roman', 12)
            run.font.italic = True
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
        run = p.add_run(stripped)
        set_run_font(run, '宋体', 'Times New Roman', 12)

        i += 1

    doc.save(str(docx_path))
    print(f"Word文档已保存: {docx_path}")
    print(f"  文件大小: {Path(docx_path).stat().st_size:,} bytes")


def main():
    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <input.md> [output.docx]")
        print("  将 Markdown 文件转换为格式化 Word 文档")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"文件不存在: {md_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix('.docx')

    create_word_doc(md_path, docx_path)


if __name__ == '__main__':
    main()
