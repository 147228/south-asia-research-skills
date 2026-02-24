# -*- coding: utf-8 -*-
"""Convert the research report Markdown to a formatted Word document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MD_PATH = Path(r"C:\Users\bisu5\Desktop\南亚研究skills\印度军政决策机制研究报告.md")
DOCX_PATH = Path(r"C:\Users\bisu5\Desktop\南亚研究skills\印度军政决策机制研究报告.docx")

# If the primary path is locked (e.g. open in Word), save to a versioned copy
import sys


def set_run_font(run, name_cn='宋体', name_en='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name_en
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from lxml import etree
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    rfonts.set(qn('w:eastAsia'), name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_formatted_text(paragraph, text, default_cn='宋体', default_size=12, default_bold=False):
    """Parse inline **bold** and *italic* markers and add runs accordingly."""
    # Pattern: **bold** or *italic* segments
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    last_end = 0
    for m in pattern.finditer(text):
        # Add text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                run = paragraph.add_run(plain)
                set_run_font(run, default_cn, 'Times New Roman', default_size, bold=default_bold)
        if m.group(2):  # **bold**
            run = paragraph.add_run(m.group(2))
            set_run_font(run, default_cn, 'Times New Roman', default_size, bold=True)
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(3))
            set_run_font(run, default_cn, 'Times New Roman', default_size, italic=True)
        last_end = m.end()
    # Remaining text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            run = paragraph.add_run(remaining)
            set_run_font(run, default_cn, 'Times New Roman', default_size, bold=default_bold)


def create_word_doc(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page setup - A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped == '---':
            # Add a small spacing paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # H1 - Report title
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            set_run_font(run, '黑体', 'Times New Roman', 18, bold=True)
            i += 1
            continue

        # H2 - Major sections
        if stripped.startswith('## ') and not stripped.startswith('### '):
            text = stripped[3:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, text, '黑体', 14, default_bold=True)
            i += 1
            continue

        # H3 - Subsections
        if stripped.startswith('### '):
            text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            add_formatted_text(p, text, '黑体', 12, default_bold=True)
            i += 1
            continue

        # Bullet list items (- item)
        if stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.37)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            bullet_run = p.add_run('  ')
            set_run_font(bullet_run, '宋体', 'Times New Roman', 12)
            add_formatted_text(p, text)
            i += 1
            continue

        # Standalone italic line (like footnote)
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            text = stripped[1:-1]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            set_run_font(run, '楷体', 'Times New Roman', 10.5, italic=True, color=(100, 100, 100))
            i += 1
            continue

        # Numbered list items (1. item)
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.37)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f'{num}. ')
            set_run_font(num_run, '宋体', 'Times New Roman', 12)
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph (with inline formatting)
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        add_formatted_text(p, stripped)
        i += 1

    doc.save(str(docx_path))
    print(f"Word document saved: {docx_path}")
    print(f"File size: {docx_path.stat().st_size:,} bytes")


if __name__ == '__main__':
    try:
        create_word_doc(MD_PATH, DOCX_PATH)
    except PermissionError:
        alt = DOCX_PATH.with_name('印度军政决策机制研究报告_v2.docx')
        print(f"Primary path locked, saving to: {alt}")
        create_word_doc(MD_PATH, alt)
